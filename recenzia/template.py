from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
import os, sys
from ctypes import *

def getRootDir(*args):
    if getattr(sys, 'frozen', False):
        application_path = os.path.abspath(os.path.dirname(sys.executable))
    else:
        application_path = os.path.dirname(__file__)
    if args:
        application_path = os.path.join(application_path, os.path.join(*args))
    application_path = application_path.replace('\\', '/')
    return application_path

def toLongName(path):
    '''
    from dos 8.3 format
    '''
    buf = create_unicode_buffer(500)
    windll.kernel32.GetLongPathNameW(path, buf, 500)
    return buf.value


def create_report(name, data, data_key_errors):
    document = Document('template.docx')

    document.add_paragraph('РЕЦЕНЗИЯ').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_paragraph('на курсовой проект {{ kat }} {{ group }} учебной группы').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_paragraph('факультета пожарной безопасности').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_paragraph('специальности 20.05.01 «Пожарная безопасность»').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_paragraph('{{ name }}').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_paragraph('Автоматизированные системы управления и связь').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_paragraph('')
    document.add_paragraph('Курсовой проект {{ conformity }}.')

    tuple_parts = ('Оформление:',
                   'Часть 1. Организация проводной связи ПСГ:',
                   'Часть 2. Организация радиосвязи ПСГ:',
                   'Часть 3. Организация сети передачи данных ПСГ:',
                   'Схемы:')
    for key in data_key_errors.keys():
        document.add_paragraph(tuple_parts[key])
        for error in data_key_errors[key]:
            paragraph = document.add_paragraph()
            paragraph_format = paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(0.2), WD_TAB_ALIGNMENT.LEFT)
            paragraph.add_run('\t' + error)
            break
        else:
            paragraph = document.add_paragraph()
            paragraph_format = paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(0.2), WD_TAB_ALIGNMENT.LEFT)
            paragraph.add_run('\tЗамечаний нет')

    document.add_paragraph('')
    document.add_paragraph('ЗАКЛЮЧЕНИЕ').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_paragraph('Курсовой проект {{ result }}.')
    document.add_paragraph('')
    document.add_paragraph('Рецензент:')
    document.add_paragraph('{{ position }}')
    document.add_paragraph('{{ rank }} внутренней службы')
    paragraph = document.add_paragraph('{{ date }}')
    paragraph_format = paragraph.paragraph_format
    tab_stops = paragraph_format.tab_stops
    tab_stop = tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT)
    paragraph.add_run('\t{{ lecturer_name }}')
    document.save('template_for_generated.docx')

    doc = DocxTemplate("template_for_generated.docx")
    context = { 'kat' : data[4],
                'group' : data[3],
                'name' : name,
                'conformity' : data[6],
                'result' : data[1],
                'position' : data[7],
                'rank' : data[8],
                'date' : data[2],
                'lecturer_name' : data[5]
                }
    doc.render(context)
    doc.save("generated.docx")

    os.startfile(os.path.join(getRootDir(), 'generated.docx'), "print")

